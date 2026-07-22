import io

from docx.shared import Pt  # noqa: F401
from django.http import HttpResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def generate_paper_pdf(paper):
    pqs = paper.paper_questions.select_related("question__topic").order_by("order")
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5 * inch, bottomMargin=0.5 * inch)
    styles = getSampleStyleSheet()
    elements = []
    school_name = paper.school.name if paper.school else ""
    if school_name:
        elements.append(Paragraph(school_name, styles["Title"]))
        elements.append(Spacer(1, 12))
    elements.append(Paragraph(paper.title, styles["Title"]))
    elements.append(Spacer(1, 12))
    meta = []
    if paper.class_name:
        meta.append(f"<b>Class:</b> {paper.class_name}")
    if paper.academic_session:
        meta.append(f"<b>Session:</b> {paper.academic_session}")
    if paper.term:
        meta.append(f"<b>Term:</b> {paper.term}")
    meta.append(f"<b>Duration:</b> {paper.duration_minutes} minutes")
    meta.append(f"<b>Total Marks:</b> {paper.total_marks}")
    elements.append(Paragraph(" | ".join(meta), styles["Normal"]))
    elements.append(Spacer(1, 12))
    if paper.instructions:
        elements.append(Paragraph(f"<b>Instructions:</b> {paper.instructions}", styles["Normal"]))
        elements.append(Spacer(1, 18))
    num = 1
    for pq in pqs:
        q = pq.question
        marks = pq.marks_override if pq.marks_override is not None else q.marks
        elements.append(Paragraph(f"<b>{num}. {q.question_text}</b> ({marks} mark{'s' if marks != 1 else ''})", styles["Normal"]))
        elements.append(Spacer(1, 6))
        if q.question_type in ("mcq", "true_false") and q.options:
            for i, opt in enumerate(q.options):
                elements.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{chr(65+i)}. {opt}", styles["Normal"]))
            elements.append(Spacer(1, 6))
        if q.question_type == "matching" and q.options:
            for i, opt in enumerate(q.options):
                elements.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{i+1}. {opt} &mdash; ________", styles["Normal"]))
            elements.append(Spacer(1, 6))
        if q.question_type == "essay":
            elements.append(Spacer(1, 40))
        num += 1
    doc.build(elements)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{paper.title.replace(" ", "_")}.pdf"'
    return response


def generate_paper_docx(paper):
    from docx import Document
    pqs = paper.paper_questions.select_related("question__topic").order_by("order")
    doc = Document()
    school_name = paper.school.name if paper.school else ""
    if school_name:
        school_heading = doc.add_heading(school_name, level=0)
        school_heading.alignment = 1  # center
    title_heading = doc.add_heading(paper.title, level=1)
    title_heading.alignment = 1  # center
    meta_parts = []
    if paper.class_name:
        meta_parts.append(f"Class: {paper.class_name}")
    if paper.academic_session:
        meta_parts.append(f"Session: {paper.academic_session}")
    if paper.term:
        meta_parts.append(f"Term: {paper.term}")
    meta_parts.append(f"Duration: {paper.duration_minutes} minutes")
    meta_parts.append(f"Total Marks: {paper.total_marks}")
    doc.add_paragraph(" | ".join(meta_parts))
    if paper.instructions:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Instructions: ")
        run.bold = True
        p.add_run(paper.instructions)
    doc.add_paragraph("")
    num = 1
    for pq in pqs:
        q = pq.question
        marks = pq.marks_override if pq.marks_override is not None else q.marks
        p = doc.add_paragraph()
        run = p.add_run(f"{num}. {q.question_text}")
        run.bold = True
        p.add_run(f"  ({marks} mark{'s' if marks != 1 else ''})")
        if q.question_type in ("mcq", "true_false") and q.options:
            for i, opt in enumerate(q.options):
                doc.add_paragraph(f"{chr(65+i)}. {opt}", style="List Bullet")
        if q.question_type == "matching" and q.options:
            for i, opt in enumerate(q.options):
                doc.add_paragraph(f"{i+1}. {opt} - ________", style="List Bullet")
        if q.question_type == "essay":
            doc.add_paragraph("")
            doc.add_paragraph("")
        doc.add_paragraph("")
        num += 1
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="{paper.title.replace(" ", "_")}.docx"'
    return response


def generate_paper_html(paper):
    pqs = paper.paper_questions.select_related("question__topic").order_by("order")
    parts = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'>",
        f"<title>{paper.title}</title>",
        "<style>body{font-family:Arial,sans-serif;margin:40px;font-size:14px;line-height:1.6;}"
        "h1{text-align:center;margin-bottom:5px;}.school-name{text-align:center;font-size:24px;font-weight:bold;margin-bottom:5px;text-transform:uppercase;letter-spacing:1px;}"
        ".meta{text-align:center;color:#555;margin-bottom:20px;}"
        ".instructions{border:1px solid #ccc;padding:10px;margin-bottom:20px;background:#f9f9f9;}"
        ".question{margin-bottom:15px;}.question-text{font-weight:bold;}"
        ".marks{color:#666;font-size:12px;}.options{margin-left:20px;margin-top:5px;}"
        "@media print{body{margin:20px;}}</style></head><body>",
    ]
    school_name = paper.school.name if paper.school else ""
    if school_name:
        parts.append(f"<div class='school-name'>{school_name}</div>")
    parts.append(f"<h1>{paper.title}</h1><div class='meta'>")
    if paper.class_name:
        parts.append(f"<strong>Class:</strong> {paper.class_name} | ")
    if paper.academic_session:
        parts.append(f"<strong>Session:</strong> {paper.academic_session} | ")
    if paper.term:
        parts.append(f"<strong>Term:</strong> {paper.term} | ")
    parts.append(f"<strong>Duration:</strong> {paper.duration_minutes} mins | <strong>Total Marks:</strong> {paper.total_marks}</div>")
    if paper.instructions:
        parts.append(f"<div class='instructions'><strong>Instructions:</strong> {paper.instructions}</div>")
    num = 1
    for pq in pqs:
        q = pq.question
        marks = pq.marks_override if pq.marks_override is not None else q.marks
        parts.append(f"<div class='question'><div class='question-text'>{num}. {q.question_text} <span class='marks'>({marks} mark{'s' if marks != 1 else ''})</span></div>")
        if q.question_type in ("mcq", "true_false") and q.options:
            parts.append("<div class='options'>")
            for i, opt in enumerate(q.options):
                parts.append(f"<div>{chr(65+i)}. {opt}</div>")
            parts.append("</div>")
        if q.question_type == "matching" and q.options:
            parts.append("<div class='options'>")
            for i, opt in enumerate(q.options):
                parts.append(f"<div>{i+1}. {opt} &mdash; ________</div>")
            parts.append("</div>")
        if q.question_type == "essay":
            parts.append("<div>&nbsp;<br>&nbsp;<br>&nbsp;</div>")
        parts.append("</div>")
        num += 1
    parts.append("<script>window.onload=function(){window.print();}</script></body></html>")
    return HttpResponse("".join(parts), content_type="text/html")
